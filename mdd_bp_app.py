import streamlit as st, pandas as pd, numpy as np, tempfile, io, os
from docx import Document
import mdd_vs_bipolar as mddbp   # import your analysis functions

# Path to Cuban norms
NORM_PATH = os.path.join(os.path.dirname(__file__), "cuban_norms.csv")

st.set_page_config(page_title="qEEG Depression vs Bipolar Demo", layout="wide")
st.title("qEEG Demo: Depression vs Bipolar")
st.write("Upload an EDF file to generate a **demo qEEG report**. "
         "⚠️ This is for demonstration only, not for clinical use.")

# --- Inputs ---
edf_file = st.file_uploader("Upload EEG (EDF format)", type=["edf"])
age = st.number_input("Patient Age", min_value=5, max_value=90, value=25)
p2p = st.number_input("Artifact rejection threshold (µV)", min_value=50, max_value=500, value=150)

if edf_file:
    with st.spinner("Processing EEG..."):
        # Save EDF temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".edf") as tmp:
            tmp.write(edf_file.read())
            edf_path = tmp.name

        # Load Cuban norms
        norms_all = pd.read_csv(NORM_PATH)
        norms_age = norms_all[norms_all["Age"] == age]
        if norms_age.empty:
            st.error(f"No norms available for Age={age}")
            st.stop()

        # --- Compute powers & Zs ---
        df_raw = mddbp.compute_raw_powers(edf_path)
        df_clean, kept = mddbp.compute_clean_powers(edf_path, p2p)
        dfz = mddbp.add_zrel_zabs(df_clean, norms_age)

        # --- Build report ---
        doc = Document()
        doc.add_heading("qEEG Report (Depression vs Bipolar) – Demo", 0)
        doc.add_paragraph(f"Age: {age}")
        doc.add_paragraph(f"Artifact rejection: {p2p} µV | Epochs kept: {kept}")
        doc.add_paragraph("Bands: Delta 1–4, Theta 4–8, Alpha 8–12 (α1:8–10, α2:10–12), Beta 12–30 Hz")

        # ---- RAW Maps ----
        doc.add_heading("RAW Relative Maps", level=1)
        raw_items = [(f"RAW – {b}", mddbp._topomap_png(df_raw, f"Rel_{b}", f"RAW – {b}")) for b in ["Delta","Theta","Alpha","Beta"]]
        mddbp.make_two_per_row_section(doc, raw_items)

        # ---- Z Maps ----
        doc.add_heading("Relative Z Maps (Cuban)", level=1)
        z_items = [(f"Z – {b}", mddbp._topomap_png(dfz, f"Zrel_{b}", f"Z – {b} (Age {age})")) for b in ["Delta","Theta","Alpha","Beta"]]
        mddbp.make_two_per_row_section(doc, z_items)

        # ---- Diagnostics ----
        doc.add_heading("Diagnostics & Markers", level=1)

        # FAA
        faa_idx, faa_tag = mddbp.compute_faa(dfz)
        faa_points = 2 if (np.isfinite(faa_idx) and faa_idx < -20) else 0
        doc.add_paragraph(f"FAA: {faa_idx:.1f} → {faa_tag} | Points={faa_points}")

        # Parietal Alpha Asymmetry
        p3_alpha = mddbp.get(dfz,"P3","Rel_Alpha")
        p4_alpha = mddbp.get(dfz,"P4","Rel_Alpha")
        if np.isfinite(p3_alpha) and np.isfinite(p4_alpha):
            diff = p3_alpha - p4_alpha
            if diff > 0.1:
                par_alpha_points = 1; asym_txt = "Left-dominant (Depression marker)"
            elif diff < -0.1:
                par_alpha_points = 0; asym_txt = "Right-dominant"
            else:
                par_alpha_points = 0; asym_txt = "Neutral"
        else:
            diff = np.nan; par_alpha_points = 0; asym_txt = "Insufficient data"
        doc.add_paragraph(f"Parietal Alpha Asymmetry (P3-P4): {diff:.3f} → {asym_txt} | Points={par_alpha_points}")

        # APF Slowing at Cz
        raw0 = mddbp.mne.io.read_raw_edf(edf_path, preload=False, verbose=False)
        cz_pick = [c for c in raw0.ch_names if mddbp.normalize_label(c) == "Cz"]
        if cz_pick:
            psd_cz, freqs_cz = mddbp._raw_psd_and_freqs(raw0, cz_pick[0], 8., 12.)
            apf_hz = freqs_cz[np.argmax(psd_cz)] if len(freqs_cz) else np.nan
            apf_points = 3 if (np.isfinite(apf_hz) and apf_hz <= 9.5) else 0
            apf_txt = "Slowed (Depression marker)" if apf_points else "Normal"
        else:
            apf_hz, apf_points, apf_txt = np.nan, 0, "Cz not found"
        doc.add_paragraph(f"APF at Cz: {apf_hz:.2f} Hz → {apf_txt} | Points={apf_points}")

        # Alpha Reduction (Bipolar)
        alpha_details, alpha_ok = mddbp.bipolar_alpha_signature(dfz)
        alpha_points = 3 if alpha_ok else 0
        doc.add_paragraph(f"Alpha reduction rule: {'MET' if alpha_ok else 'NOT MET'} | Points={alpha_points}")

        # Theta Right-lower chain (Bipolar)
        sites, hits, theta_ok = mddbp.bipolar_theta_signature(dfz)
        theta_points = 2 if theta_ok else 0
        doc.add_paragraph(f"Theta right-lower chain: {'MET' if theta_ok else 'NOT MET'} | Points={theta_points}")

        # Parietal Beta Bilateral (Bipolar)
        p3_beta = mddbp.get(dfz, "P3", "Zrel_Beta")
        p4_beta = mddbp.get(dfz, "P4", "Zrel_Beta")
        beta_points = 1 if (np.isfinite(p3_beta) and np.isfinite(p4_beta) and p3_beta>=1.0 and p4_beta>=1.0) else 0
        doc.add_paragraph(f"Parietal Beta Bilateral: P3={p3_beta:.2f}, P4={p4_beta:.2f} | Points={beta_points}")

        # ---- Final Scoring ----
        dep_score = faa_points + par_alpha_points + apf_points
        bp_score  = alpha_points + theta_points + beta_points
        total = dep_score + bp_score
        if total > 0:
            p_dep = dep_score / total; p_bp = bp_score / total
        else:
            p_dep = p_bp = 0.5
        final_call = "Depression-leaning" if dep_score>bp_score else ("Bipolar-leaning" if bp_score>dep_score else "Ambiguous")

        doc.add_heading("Probability Scores", level=1)
        doc.add_paragraph(f"P(Depression) = {p_dep*100:.1f}%")
        doc.add_paragraph(f"P(Bipolar)    = {p_bp*100:.1f}%")
        doc.add_heading(f"Final Call: {final_call}", level=1)

        # Save buffer
        buffer = io.BytesIO(); doc.save(buffer)

    st.success(f"Report ready ✅ Final Call: {final_call}")
    st.download_button("Download Report", buffer.getvalue(),
                       file_name="qEEG_MDDvsBipolar_Report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
