#!/usr/bin/env python3
# qEEG Report: ADHD Subtyping (Theta Excess & TBR rules)
# Implements research-based ADHD subtyping with updated thresholds:
# 1. Theta Excess rule: Z_Theta > 1 (changed from >= 1.5) in at least two frontal electrodes
# 2. TBR (Theta/Beta Ratio) rule: Age-dependent thresholds at Cz

import io, os
import numpy as np
import pandas as pd
import mne
from scipy.interpolate import griddata
from scipy.signal import welch as scipy_welch
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Circle
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox

# -------------------- Config --------------------
BANDS = {"Delta": (1,4), "Theta": (4,8), "Alpha": (8,12), "Beta": (12,30)}
TOTAL  = (1,30)
ALPHA1 = (8,10)
ALPHA2 = (10,12)

# 10-20 scalp positions for simple topomap (approximate 2D)
POS = {
    "Fp1":(-0.5,0.9),"Fp2":(0.5,0.9),
    "F7":(-0.9,0.5),"F3":(-0.4,0.5),"Fz":(0.0,0.55),"F4":(0.4,0.5),"F8":(0.9,0.5),
    "T7":(-1.0,0.0),"C3":(-0.5,0.0),"Cz":(0.0,0.0),"C4":(0.5,0.0),"T8":(1.0,0.0),
    "P7":(-0.9,-0.5),"P3":(-0.4,-0.5),"Pz":(0.0,-0.55),"P4":(0.4,-0.5),"P8":(0.9,-0.5),
    "O1":(-0.5,-0.9),"O2":(0.5,-0.9)
}

# Regions for ADHD subtyping rules
REGIONS = {
    "Frontal":  ["Fp1","Fp2","F3","F4","Fz"],
    "Temporal": ["T7","T8","P7","P8"],
    "Central":  ["C3","C4","Cz"],
    "Parietal": ["P3","P4","Pz"],
    "Posterior":["P3","P4","Pz","O1","O2"]
}

# -------------------- Label normalization --------------------
def normalize_label(ch: str) -> str:
    """Normalize EDF channel labels to Cuban/10-20 standard names."""
    ch = ch.strip().upper()

    # Remove common prefixes like "EEG " or "EEG."
    if ch.startswith("EEG "):
        ch = ch.replace("EEG ", "")
    if ch.startswith("EEG."):
        ch = ch.replace("EEG.", "")

    # Trim common suffixes for references
    for suf in ("-A1","-A2","-LE","-RE","-M1","-M2","-AVG","-REF"):
        if ch.endswith(suf):
            ch = ch.split("-")[0]
            break

    # Map 10-20 alternatives (T3/T4/T5/T6 -> T7/T8/P7/P8)
    mapping = {
        "FP1":"Fp1","FP2":"Fp2",
        "F7":"F7","F3":"F3","FZ":"Fz","F4":"F4","F8":"F8",
        "T3":"T7","T4":"T8","T5":"P7","T6":"P8",
        "C3":"C3","CZ":"Cz","C4":"C4",
        "P3":"P3","PZ":"Pz","P4":"P4",
        "O1":"O1","O2":"O2"
    }
    return mapping.get(ch, ch.title())

def check_channel_mapping(raw, norms):
    cuban_channels = set(norms["Channel"].unique())
    unmatched = []
    for ch in raw.ch_names:
        mapped = normalize_label(ch)
        if mapped not in cuban_channels:
            unmatched.append((ch, mapped))
    return unmatched

# -------------------- PSD helpers --------------------
def _integ(psd, freqs, lo, hi):
    m = (freqs>=lo)&(freqs<hi)
    return float(np.trapz(psd[m], freqs[m]))

def _raw_psd_and_freqs(raw, ch_name, fmin=1., fmax=30.):
    sf = raw.info["sfreq"]
    data = raw.get_data(picks=[ch_name])[0]
    nper = int(round(2*sf)); nover = nper//2
    f, pxx = scipy_welch(data, fs=sf, nperseg=nper, noverlap=nover, scaling='density')
    sel = (f>=fmin)&(f<=fmax)
    return pxx[sel], f[sel]

def _epochs_psd_and_freqs(epochs, fmin=1., fmax=30.):
    # Newer MNE API
    if hasattr(epochs, "compute_psd"):
        psd_obj = epochs.compute_psd(fmin=fmin, fmax=fmax, method="welch", verbose=False)
        psds, freqs = psd_obj.get_data(return_freqs=True)
        return psds, freqs
    # Fallback
    data = epochs.get_data()
    sf = epochs.info["sfreq"]
    nper = int(round(2*sf)); nover = nper//2
    psd_list = []
    for e in range(data.shape[0]):
        chs = []
        for c in range(data.shape[1]):
            f, pxx = scipy_welch(data[e, c, :], fs=sf, nperseg=nper, noverlap=nover, scaling='density')
            sel = (f>=fmin)&(f<=fmax)
            chs.append(pxx[sel]); freqs = f[sel]
        psd_list.append(np.vstack(chs))
    psds = np.stack(psd_list, axis=0)
    return psds, freqs

# -------------------- Power computations --------------------
def compute_raw_powers(edf_path: str) -> pd.DataFrame:
    """RAW (single-trace) relative/absolute band power per channel (for raw maps)."""
    raw = mne.io.read_raw_edf(edf_path, preload=True, verbose=False)
    rows = []
    for ch in raw.ch_names:
        psd, freqs = _raw_psd_and_freqs(raw, ch, 1., 30.)
        total = _integ(psd, freqs, *TOTAL)
        band_abs = {b: _integ(psd, freqs, lo,hi) for b,(lo,hi) in BANDS.items()}
        rel = {b: (band_abs[b]/total if total>0 else np.nan) for b in BANDS}
        a1 = _integ(psd, freqs, *ALPHA1)
        a2 = _integ(psd, freqs, *ALPHA2)
        rows.append({
            "EDF_Channel": ch, "Cuban_Channel": normalize_label(ch),
            **{f"Abs_{b}": band_abs[b] for b in BANDS},
            **{f"Rel_{b}": rel[b] for b in BANDS},
            "Abs_Alpha1": a1, "Abs_Alpha2": a2
        })
    df = pd.DataFrame(rows)
    df["Abs_Alpha"] = df["Abs_Alpha1"] + df["Abs_Alpha2"]
    return df

def compute_clean_powers(edf_path: str, p2p_uv: float = 150.0) -> tuple[pd.DataFrame,int]:
    """CLEAN (2s epochs, artifact-rejected) average band power per channel (for Z maps)."""
    raw = mne.io.read_raw_edf(edf_path, preload=True, verbose=False)
    events = mne.make_fixed_length_events(raw, duration=2.0)
    epochs = mne.Epochs(raw, events=events, tmin=0., tmax=2.0,
                        baseline=None, preload=True, verbose=False)
    epochs.drop_bad(reject=dict(eeg=p2p_uv*1e-6), verbose=False)
    if len(epochs)==0:
        # fallback, a bit looser
        epochs = mne.Epochs(raw, events=events, tmin=0., tmax=2.0,
                            baseline=None, preload=True, verbose=False)
        epochs.drop_bad(reject=dict(eeg=200e-6), verbose=False)

    psds, freqs = _epochs_psd_and_freqs(epochs, fmin=1., fmax=30.)
    psd_mean = psds.mean(axis=0)  # (n_channels, n_freqs)

    def integ(s, lo, hi):
        m = (freqs>=lo)&(freqs<hi)
        return float(np.trapz(s[m], freqs[m]))

    rows=[]
    for i,ch in enumerate(epochs.ch_names):
        s = psd_mean[i]
        total = integ(s, *TOTAL)
        band_abs = {b: integ(s, lo,hi) for b,(lo,hi) in BANDS.items()}
        rel = {b: (band_abs[b]/total if total>0 else np.nan) for b in BANDS}
        a1 = integ(s, *ALPHA1); a2 = integ(s, *ALPHA2)
        rows.append({"EDF_Channel": ch, "Cuban_Channel": normalize_label(ch),
                     **{f"Abs_{b}": band_abs[b] for b in BANDS},
                     **{f"Rel_{b}": rel[b] for b in BANDS},
                     "Abs_Alpha1": a1, "Abs_Alpha2": a2})
    df = pd.DataFrame(rows)
    df["Abs_Alpha"] = df["Abs_Alpha1"] + df["Abs_Alpha2"]
    return df, len(epochs)

# -------------------- Z-score helper --------------------
def add_zrel_zabs(df: pd.DataFrame, norms_age: pd.DataFrame) -> pd.DataFrame:
    """Add Cuban relative/absolute Z scores (assumes norms CSV with Age, Channel, Band, means/SDs)."""
    def z_rel(p,m,s): return (np.log10(p)-m)/s if (p>0 and np.isfinite(p)) else np.nan
    def z_abs(v,m,s): return (np.log10(v)+12.0-m)/s if (v>0 and np.isfinite(v)) else np.nan
    out = df.copy()
    for b in BANDS:
        zr, za = [], []
        for _,r in out.iterrows():
            ch = r["Cuban_Channel"]
            nr = norms_age[(norms_age["Channel"]==ch)&(norms_age["Band"]==b)]
            if nr.empty:
                zr.append(np.nan); za.append(np.nan)
            else:
                nr = nr.iloc[0]
                zr.append(z_rel(r[f"Rel_{b}"], nr["RelMean"], nr["RelSD"]))
                za.append(z_abs(r[f"Abs_{b}"], nr["AbsMean_log10"], nr["AbsSD_log10"]))
        out[f"Zrel_{b}"]=zr; out[f"Zabs_{b}"]=za
    return out

# -------------------- Plotting --------------------
def _topomap_png(df: pd.DataFrame, key: str, title: str) -> bytes:
    xs, ys, zs = [], [], []
    for ch,(x,y) in POS.items():
        row = df[df["Cuban_Channel"]==ch]
        if row.empty: continue
        v = row.iloc[0].get(key, np.nan)
        if np.isfinite(v):
            xs.append(x); ys.append(y); zs.append(float(v))
    fig = plt.figure(figsize=(4.6,3.9))
    ax = fig.add_subplot(111, aspect="equal")
    if len(xs)==0:
        ax.text(0.5,0.5,"No data", ha="center", va="center")
    else:
        xs,ys,zs = np.array(xs), np.array(ys), np.array(zs)
        grid_x, grid_y = np.mgrid[-1:1:220j, -1:1:220j]
        pts = np.vstack([xs,ys]).T
        grid_z = griddata(pts, zs, (grid_x, grid_y), method="cubic")
        r = np.sqrt(grid_x**2 + grid_y**2); grid_z[r>1]=np.nan
        if key.startswith("Z"):
            levels = np.linspace(-3,3,13)
            cf = ax.contourf(grid_x, grid_y, grid_z, levels=levels, cmap="RdBu_r", vmin=-3, vmax=3)
        else:
            cf = ax.contourf(grid_x, grid_y, grid_z, levels=12)
        plt.colorbar(cf, ax=ax, shrink=0.75)
    ax.add_patch(Circle((0,0),1.0, fill=False, lw=2))
    ax.scatter([p[0] for p in POS.values()],[p[1] for p in POS.values()], c="k", s=12)
    ax.set_xticks([]); ax.set_yticks([]); ax.set_title(title)
    bio=io.BytesIO(); fig.savefig(bio, format="png", dpi=160, bbox_inches="tight"); plt.close(fig)
    return bio.getvalue()

def make_two_per_row_section(doc: Document, items: list[tuple[str, bytes]]):
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(0, len(items), 2):
        row_cells = table.add_row().cells
        for j in range(2):
            idx = i + j
            if idx >= len(items): break
            cap, png = items[idx]
            cell = row_cells[j]
            para = cell.paragraphs[0]; run = para.add_run()
            run.add_picture(io.BytesIO(png), width=Inches(2.8))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cap_p = cell.add_paragraph(cap); cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

# -------------------- Utilities --------------------
def get(df, ch, key):
    """Get value for a specific channel and key from dataframe."""
    r = df[df["Cuban_Channel"]==ch]
    return float(r.iloc[0][key]) if not r.empty and key in r else np.nan

# -------------------- ADHD Subtyping Functions --------------------
def theta_excess_rule(dfz: pd.DataFrame):
    """
    Theta Excess rule: Z_Theta > 1 (changed from >= 1.5) in at least two frontal electrodes.
    
    Returns:
        sites (dict): Z_Theta values for each frontal electrode
        hits (dict): Boolean indicators for each frontal electrode meeting threshold
        is_positive (bool): True if rule is met (>=2 frontal sites with Z_Theta > 1)
    """
    frontal_electrodes = REGIONS["Frontal"]
    sites = {ch: get(dfz, ch, "Zrel_Theta") for ch in frontal_electrodes}
    hits = {ch: (np.isfinite(v) and v > 1.0) for ch, v in sites.items()}
    count = sum(1 for h in hits.values() if h)
    is_positive = (count >= 2)
    
    return sites, hits, is_positive

def compute_tbr_at_cz(dfz: pd.DataFrame):
    """
    Compute Theta/Beta Ratio (TBR) at Cz electrode.
    
    Returns:
        tbr_value (float): TBR value at Cz (nan if insufficient data)
        theta_cz (float): Theta power at Cz
        beta_cz (float): Beta power at Cz
    """
    # Use relative power values for TBR calculation
    theta_cz = get(dfz, "Cz", "Rel_Theta")
    beta_cz = get(dfz, "Cz", "Rel_Beta")
    
    if np.isfinite(theta_cz) and np.isfinite(beta_cz) and beta_cz > 0:
        tbr_value = theta_cz / beta_cz
    else:
        tbr_value = np.nan
    
    return tbr_value, theta_cz, beta_cz

def high_tbr_rule(dfz: pd.DataFrame, age: int):
    """
    TBR (Theta/Beta Ratio) rule with age-dependent thresholds:
    - For children and adolescents (age < 19): High TBR subtype if TBR > 4 at Cz
    - For adults (age >= 19): High TBR subtype if TBR > 3.0 at Cz
    
    Args:
        dfz: DataFrame with Z-scored EEG data
        age: Patient age in years
    
    Returns:
        tbr_value (float): TBR value at Cz
        threshold (float): Age-appropriate threshold used
        is_high (bool): True if TBR exceeds threshold
        age_category (str): Age category description
    """
    tbr_value, theta_cz, beta_cz = compute_tbr_at_cz(dfz)
    
    if age < 19:
        threshold = 4.0
        age_category = "Child/Adolescent"
    else:
        threshold = 3.0
        age_category = "Adult"
    
    is_high = (np.isfinite(tbr_value) and tbr_value > threshold)
    
    return tbr_value, threshold, is_high, age_category

def classify_adhd_subtypes(dfz: pd.DataFrame, age: int):
    """
    Main ADHD subtyping classification function.
    
    Args:
        dfz: DataFrame with Z-scored EEG data
        age: Patient age in years
    
    Returns:
        results (dict): Complete classification results
    """
    # Apply Theta Excess rule
    theta_sites, theta_hits, theta_excess = theta_excess_rule(dfz)
    
    # Apply TBR rule
    tbr_value, tbr_threshold, high_tbr, age_category = high_tbr_rule(dfz, age)
    
    # Determine ADHD subtype
    if theta_excess and high_tbr:
        subtype = "ADHD Combined Type (Theta Excess + High TBR)"
        subtype_score = 2
    elif theta_excess:
        subtype = "ADHD Inattentive Type (Theta Excess)"
        subtype_score = 1
    elif high_tbr:
        subtype = "ADHD Hyperactive/Impulsive Type (High TBR)"
        subtype_score = 1
    else:
        subtype = "No ADHD EEG markers detected"
        subtype_score = 0
    
    results = {
        "theta_excess": {
            "sites": theta_sites,
            "hits": theta_hits,
            "is_positive": theta_excess,
            "count": sum(1 for h in theta_hits.values() if h)
        },
        "tbr": {
            "value": tbr_value,
            "threshold": tbr_threshold,
            "is_high": high_tbr,
            "age_category": age_category
        },
        "classification": {
            "subtype": subtype,
            "score": subtype_score
        }
    }
    
    return results

# -------------------- Main --------------------
def main():
    root = tk.Tk(); root.withdraw()
    try:
        edf_path = filedialog.askopenfilename(title="Select EDF file", filetypes=[("EDF","*.edf"),("All","*.*")])
        if not edf_path: raise RuntimeError("EDF not selected.")
        norms_path = filedialog.askopenfilename(title="Select Cuban norms CSV", filetypes=[("CSV","*.csv"),("All","*.*")])
        if not norms_path: raise RuntimeError("Norms CSV not selected.")

        raw0 = mne.io.read_raw_edf(edf_path, preload=False, verbose=False)
        norms_all = pd.read_csv(norms_path)
        unmatched = check_channel_mapping(raw0, norms_all)

        age = simpledialog.askinteger("Patient Age","Enter patient age (Cuban norms must include this age):",minvalue=1,maxvalue=120)
        if age is None: raise RuntimeError("Age not provided.")
        p2p = simpledialog.askfloat("Artifact threshold (µV)","Drop epochs with p2p > µV (default 150):",initialvalue=150.0,minvalue=50.0,maxvalue=500.0)
        if p2p is None: p2p = 150.0
        out_docx = filedialog.asksaveasfilename(title="Save Word report as...",defaultextension=".docx",filetypes=[("Word Document","*.docx")])
        if not out_docx: raise RuntimeError("Output path not selected.")

        # Compute powers and Zs
        df_raw = compute_raw_powers(edf_path)                 # RAW relative maps
        df_clean, kept = compute_clean_powers(edf_path, p2p)  # CLEAN powers for Z maps
        norms_age = norms_all[norms_all["Age"]==age].copy()
        if norms_age.empty:
            raise RuntimeError(f"No norms found for Age={age}.")
        dfz = add_zrel_zabs(df_clean, norms_age)

        # ADHD Classification
        adhd_results = classify_adhd_subtypes(dfz, age)

        # Build report
        doc = Document()
        doc.add_heading("qEEG Report: ADHD Subtyping", level=0)
        doc.add_paragraph(f"EDF: {os.path.basename(edf_path)}")
        doc.add_paragraph(f"Norms: {os.path.basename(norms_path)}  |  Age: {age}")
        doc.add_paragraph(f"Artifact rejection: 2-s epochs, drop p2p > {p2p:.0f} µV. Kept epochs: {kept}")
        doc.add_paragraph("Band definitions: Delta 1–4, Theta 4–8, Alpha 8–12 (Alpha1 8–10, Alpha2 10–12), Beta 12–30 Hz.")

        if unmatched:
            doc.add_paragraph("Channels not found in Cuban norms (skipped in Z-maps):")
            for orig, mapped in unmatched:
                doc.add_paragraph(f"  EDF: {orig}  ->  Normalized: {mapped}")
        else:
            doc.add_paragraph("All EDF channels matched Cuban norms.")

        # RAW Relative maps
        doc.add_heading("RAW Relative Maps", level=1)
        raw_items = [(f"RAW Relative – {b}", _topomap_png(df_raw, f"Rel_{b}", f"RAW Relative – {b}"))
                     for b in ["Delta","Theta","Alpha","Beta"]]
        make_two_per_row_section(doc, raw_items)

        # CLEAN Z-Relative maps
        doc.add_heading("Relative Z Maps (Cuban, CLEAN)", level=1)
        z_items = [(f"Z (Cuban) – {b}", _topomap_png(dfz, f"Zrel_{b}", f"Z – {b} (Age {age})"))
                   for b in ["Delta","Theta","Alpha","Beta"]]
        make_two_per_row_section(doc, z_items)

        # ---------------- ADHD Subtyping Analysis ----------------
        doc.add_heading("ADHD Subtyping Analysis", level=1)

        # Theta Excess Rule
        doc.add_heading("1. Theta Excess Rule", level=2)
        theta_res = adhd_results["theta_excess"]
        doc.add_paragraph(f"Rule: Z_Theta > 1.0 in at least 2 frontal electrodes")
        doc.add_paragraph(f"Frontal electrodes analyzed: {', '.join(REGIONS['Frontal'])}")
        
        frontal_details = []
        for ch in REGIONS["Frontal"]:
            z_val = theta_res["sites"][ch]
            hit = theta_res["hits"][ch]
            if np.isfinite(z_val):
                frontal_details.append(f"{ch}: {z_val:.2f} {'✓' if hit else '✗'}")
            else:
                frontal_details.append(f"{ch}: N/A")
        
        doc.add_paragraph(f"Z_Theta values: {' | '.join(frontal_details)}")
        doc.add_paragraph(f"Electrodes meeting threshold: {theta_res['count']}/5")
        doc.add_paragraph(f"Result: {'POSITIVE' if theta_res['is_positive'] else 'NEGATIVE'} (Theta Excess)")

        # TBR Rule
        doc.add_heading("2. Theta/Beta Ratio (TBR) Rule", level=2)
        tbr_res = adhd_results["tbr"]
        doc.add_paragraph(f"Age category: {tbr_res['age_category']} (Age: {age})")
        doc.add_paragraph(f"Threshold: TBR > {tbr_res['threshold']:.1f} at Cz")
        
        if np.isfinite(tbr_res["value"]):
            doc.add_paragraph(f"TBR at Cz: {tbr_res['value']:.2f}")
        else:
            doc.add_paragraph("TBR at Cz: N/A (insufficient data)")
        
        doc.add_paragraph(f"Result: {'POSITIVE' if tbr_res['is_high'] else 'NEGATIVE'} (High TBR)")

        # Final Classification
        doc.add_heading("ADHD Subtype Classification", level=2)
        classification = adhd_results["classification"]
        doc.add_paragraph(f"Final Classification: {classification['subtype']}")
        doc.add_paragraph(f"EEG Marker Score: {classification['score']}/2")

        # Interpretation
        doc.add_heading("Clinical Interpretation", level=2)
        doc.add_paragraph("Updated ADHD subtyping rules (Research-based):")
        doc.add_paragraph("• Theta Excess: Changed threshold from Z ≥ 1.5 to Z > 1.0 for increased sensitivity")
        doc.add_paragraph("• TBR Thresholds: Age-dependent criteria reflecting developmental changes")
        doc.add_paragraph("  - Children/Adolescents (<19 years): TBR > 4.0")
        doc.add_paragraph("  - Adults (≥19 years): TBR > 3.0")
        doc.add_paragraph("")
        doc.add_paragraph("Note: These EEG markers are research-based indicators and should be interpreted")
        doc.add_paragraph("in conjunction with comprehensive clinical assessment.")

        # Save
        doc.save(out_docx)
        messagebox.showinfo("Done", f"ADHD Subtyping Report saved:\n{out_docx}")

    except Exception as e:
        messagebox.showerror("Error", str(e))
        raise

if __name__ == "__main__":
    main()