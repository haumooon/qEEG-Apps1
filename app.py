import streamlit as st
import pandas as pd
import numpy as np
import tempfile, io, os
from docx import Document
from ocd import (
    compute_raw_powers, compute_clean_powers, add_zrel_zabs,
    compute_faa, region_mean, _topomap_png,
    REGIONS
)

# Path to Cuban norms (must be in same folder as app.py)
NORM_PATH = os.path.join(os.path.dirname(__file__), "cuban_norms.csv")

st.set_page_config(page_title="qEEG OCD vs Anxiety Demo", layout="wide")
st.title("qEEG Demo: OCD vs Anxiety/Panic")
st.write("Upload an EDF file to generate a **demo qEEG report**. "
         "⚠️ This is for demonstration only, not for clinical use.")

# --- Inputs ---
edf_file = st.file_uploader("Upload EEG (EDF format)", type=["edf"])
age = st.number_input("Patient Age", min_value=5, max_value=90, value=25)
p2p = st.number_input("Artifact rejection threshold (µV)", min_value=50, max_value=500, value=150)

if edf_file:
    with st.spinner("Processing EEG..."):
        # Save EDF to a temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".edf") as tmp_edf:
            tmp_edf.write(edf_file.read())
            edf_path = tmp_edf.name

        # Load Cuban norms
        norms_all = pd.read_csv(NORM_PATH)
        norms_age = norms_all[norms_all["Age"] == age]
        if norms_age.empty:
            st.error(f"No norms available for Age={age}")
            st.stop()

        # --- Run analysis ---
        df_raw = compute_raw_powers(edf_path)                  # RAW relative maps
        df_clean, kept = compute_clean_powers(edf_path, p2p)   # CLEAN powers
        dfz = add_zrel_zabs(df_clean, norms_age)

        # --- Region scores ---
        zF_th = region_mean(dfz, "Zrel_Theta", REGIONS["Frontal"])
        zT_th = region_mean(dfz, "Zrel_Theta", REGIONS["Temporal"])
        zFT_th = np.nanmean([v for v in [zF_th, zT_th] if np.isfinite(v)]) if (np.isfinite(zF_th) or np.isfinite(zT_th)) else np.nan
        zF_be = region_mean(dfz, "Zrel_Beta", REGIONS["Frontal"])
        zF_al = region_mean(dfz, "Zrel_Alpha", REGIONS["Frontal"])
        zC_be = region_mean(dfz, "Zrel_Beta", REGIONS["Central"])
        zP_be = region_mean(dfz, "Zrel_Beta", REGIONS["Parietal"])
        zPost_be = region_mean(dfz, "Zrel_Beta", REGIONS["Posterior"])
        beta_diffuse = (np.nanmean([zC_be, zP_be]) if (np.isfinite(zC_be) or np.isfinite(zP_be)) else np.nan) - (zF_be if np.isfinite(zF_be) else 0.0)

        faa_idx, faa_tag = compute_faa(dfz)

        # --- Scoring rules ---
        ocd_theta_pts      = 3 if (np.isfinite(zFT_th) and zFT_th >= 0.5) else 0
        ocd_alpha_red_pts  = 1 if (np.isfinite(zF_al) and zF_al <= -0.5) else 0
        ocd_beta_pts       = 2 if (np.isfinite(zF_be) and zF_be >= 0.5) else 0

        anx_faa_pts            = 2 if (np.isfinite(faa_idx) and faa_idx < -20) else 0
        anx_beta_diffuse_pts   = 2 if (np.isfinite(beta_diffuse) and beta_diffuse >= 0.5) else 0
        anx_posterior_beta_pts = 1 if (np.isfinite(zPost_be) and zPost_be >= 0.5) else 0
        anx_theta_not_high_pts = 3 if (np.isfinite(zFT_th) and zFT_th < 0.5) else 0

        ocd_score = ocd_theta_pts + ocd_alpha_red_pts + ocd_beta_pts
        anx_score = anx_faa_pts + anx_beta_diffuse_pts + anx_posterior_beta_pts + anx_theta_not_high_pts
        total = ocd_score + anx_score
        if total > 0:
            p_ocd = ocd_score / total
            p_anx = anx_score / total
        else:
            p_ocd = p_anx = 0.5

        final_call = "OCD-leaning" if ocd_score > anx_score else ("Anxiety/Panic-leaning" if anx_score > ocd_score else "Ambiguous")

        # --- Build Word report ---
        doc = Document()
        doc.add_heading("qEEG Report (OCD vs Anxiety/Panic) – Demo", 0)
        doc.add_paragraph(f"Age: {age}")
        doc.add_paragraph(f"Artifact rejection: {p2p} µV | Epochs kept: {kept}")
        doc.add_paragraph("Band definitions: Delta 1–4, Theta 4–8, Alpha 8–12, Beta 12–30 Hz")

        doc.add_heading("Markers", level=1)
        doc.add_paragraph(f"Frontal/Temporal Theta: {zFT_th:.2f} → Points {ocd_theta_pts}")
        doc.add_paragraph(f"Frontal Beta: {zF_be:.2f} → Points {ocd_beta_pts}")
        doc.add_paragraph(f"Frontal Alpha: {zF_al:.2f} → Points {ocd_alpha_red_pts}")
        doc.add_paragraph(f"FAA: {faa_idx:.1f} → {faa_tag} → Points {anx_faa_pts}")
        doc.add_paragraph(f"Diffuse Beta: {beta_diffuse:.2f} → Points {anx_beta_diffuse_pts}")
        doc.add_paragraph(f"Posterior Beta: {zPost_be:.2f} → Points {anx_posterior_beta_pts}")
        doc.add_paragraph(f"Theta not high (<0.5): {'Yes' if anx_theta_not_high_pts==3 else 'No'} → Points {anx_theta_not_high_pts}")

        doc.add_heading("Probability Scores", level=1)
        doc.add_paragraph(f"P(OCD) = {p_ocd*100:.1f}%")
        doc.add_paragraph(f"P(Anxiety/Panic) = {p_anx*100:.1f}%")
        doc.add_heading(f"Final Call: {final_call}", level=1)

        # --- Add 2×2 grid of maps ---
        doc.add_heading("Z-Relative Maps (Cuban Norms)", level=1)
        maps = []
        for band in ["Delta", "Theta", "Alpha", "Beta"]:
            img_bytes = _topomap_png(dfz, f"Zrel_{band}", f"Z – {band} (Age {age})")
            maps.append((band, img_bytes))

        table = doc.add_table(rows=2, cols=2)
        row, col = 0, 0
        for band, img in maps:
            cell = table.cell(row, col)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(img), width=st.session_state.get("img_width", 2.5))
            cell.add_paragraph(band).alignment = 1
            col += 1
            if col == 2:
                row += 1
                col = 0

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)

    st.success(f"Report ready ✅ Final Call: {final_call}")
    st.download_button(
        "Download Demo Report",
        buffer.getvalue(),
        file_name="qEEG_Demo_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
