import streamlit as st

st.title("qEEG Apps1")

# Download button for Farsi Report (before any analysis)
try:
    with open("Farsi Report.docx", "rb") as file:
        doc_bytes = file.read()
    st.download_button(
        label="Download Farsi Report",
        data=doc_bytes,
        file_name="Farsi Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
except Exception as e:
    st.warning("Farsi Report.docx not found. Please add it to the app folder.")

import io
import os
import numpy as np
import pandas as pd
import mne
from scipy.interpolate import griddata
from scipy.signal import welch as scipy_welch
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Circle
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# -------------------- Config --------------------
BANDS = {"Delta": (1,4), "Theta": (4,8), "Alpha": (8,12), "Beta": (12,30)}
TOTAL  = (1,30)
ALPHA1 = (8,10)
ALPHA2 = (10,12)

POS = {
    "Fp1":(-0.5,0.9),"Fp2":(0.5,0.9),
    "F7":(-0.9,0.5),"F3":(-0.4,0.5),"Fz":(0.0,0.55),"F4":(0.4,0.5),"F8":(0.9,0.5),
    "T7":(-1.0,0.0),"C3":(-0.5,0.0),"Cz":(0.0,0.0),"C4":(0.5,0.0),"T8":(1.0,0.0),
    "P7":(-0.9,-0.5),"P3":(-0.4,-0.5),"Pz":(0.0,-0.55),"P4":(0.4,-0.5),"P8":(0.9,-0.5),
    "O1":(-0.5,-0.9),"O2":(0.5,-0.9)
}

REGIONS = {
    "Central":  ["C3","C4","Cz"],
    "Parietal": ["P3","P4","P7","P8","Pz"],
    "Occipital":["O1","O2"]
}

# -------------------- Label normalization --------------------
def normalize_label(ch: str) -> str:
    ch = ch.strip().upper()
    if ch.startswith("EEG "):
        ch = ch.replace("EEG ", "")
    if ch.startswith("EEG."):
        ch = ch.replace("EEG.", "")
    for suf in ("-A1","-A2","-LE","-RE","-M1","-M2","-AVG","-REF"):
        if ch.endswith(suf):
            ch = ch.split("-")[0]
            break
    mapping = {
        "FP1":"Fp1", "FP2":"Fp2",
        "F7":"F7", "F3":"F3", "FZ":"Fz", "F4":"F4", "F8":"F8",
        "T3":"T7", "T4":"T8", "T5":"P7", "T6":"P8",
        "C3":"C3", "CZ":"Cz", "C4":"C4",
        "P3":"P3", "PZ":"Pz", "P4":"P4",
        "O1":"O1", "O2":"O2"
    }
    return mapping.get(ch, ch.title())

def check_channel_mapping(raw, norms):
    cuban_channels = set(norms["Channel"].unique())
    unmatched = []
    for ch in raw.ch_names:
        mapped = normalize_label(ch)
        if mapped not in cuban_channels:
            unmatched.append((ch, mapped))
    return unmatched

# -------------------- PSD helpers --------------------
def _integ(psd, freqs, lo, hi):
    m = (freqs>=lo)&(freqs<hi)
    return float(np.trapz(psd[m], freqs[m]))

def _raw_psd_and_freqs(raw, ch_name, fmin=1., fmax=30.):
    sf = raw.info["sfreq"]
    data = raw.get_data(picks=[ch_name])[0]
    nper = int(round(2*sf)); nover = nper//2
    f, pxx = scipy_welch(data, fs=sf, nperseg=nper, noverlap=nover, scaling='density')
    sel = (f>=fmin)&(f<=fmax)
    return pxx[sel], f[sel]

def _epochs_psd_and_freqs(epochs, fmin=1., fmax=30.):
    if hasattr(epochs, "compute_psd"):
        psd_obj = epochs.compute_psd(fmin=fmin, fmax=fmax, method="welch", verbose=False)
        psds, freqs = psd_obj.get_data(return_freqs=True)
        return psds, freqs
    data = epochs.get_data()
    sf = epochs.info["sfreq"]
    nper = int(round(2*sf))
    nover = nper//2
    psd_list = []
    for e in range(data.shape[0]):
        chs = []
        for c in range(data.shape[1]):
            f, pxx = scipy_welch(data[e, c, :], fs=sf, nperseg=nper, noverlap=nover, scaling='density')
            sel = (f>=fmin)&(f<=fmax)
            chs.append(pxx[sel])
            freqs = f[sel]
        psd_list.append(np.vstack(chs))
    psds = np.stack(psd_list, axis=0)
    return psds, freqs

# -------------------- Power computations --------------------
def compute_raw_powers(edf_path: str) -> pd.DataFrame:
    raw = mne.io.read_raw_edf(edf_path, preload=True, verbose=False)
    rows = []
    for ch in raw.ch_names:
        psd, freqs = _raw_psd_and_freqs(raw, ch, 1., 30.)
        total = _integ(psd, freqs, *TOTAL)
        band_abs = {b: _integ(psd, freqs, lo,hi) for b,(lo,hi) in BANDS.items()}
        rel = {b: (band_abs[b]/total if total>0 else np.nan) for b in BANDS}
        a1 = _integ(psd, freqs, *ALPHA1)
        a2 = _integ(psd, freqs, *ALPHA2)
        rows.append({
            "EDF_Channel": ch, "Cuban_Channel": normalize_label(ch),
            **{f"Abs_{b}": band_abs[b] for b in BANDS},
            **{f"Rel_{b}": rel[b] for b in BANDS},
            "Abs_Alpha1": a1, "Abs_Alpha2": a2
        })
    df = pd.DataFrame(rows)
    df["Abs_Alpha"] = df["Abs_Alpha1"] + df["Abs_Alpha2"]
    return df

def compute_clean_powers(edf_path: str, p2p_uv: float = 150.0) -> tuple[pd.DataFrame,int]:
    raw = mne.io.read_raw_edf(edf_path, preload=True, verbose=False)
    events = mne.make_fixed_length_events(raw, duration=2.0)
    epochs = mne.Epochs(raw, events=events, tmin=0., tmax=2.0,
                        baseline=None, preload=True, verbose=False)
    epochs.drop_bad(reject=dict(eeg=p2p_uv*1e-6), verbose=False)
    if len(epochs)==0:
        epochs = mne.Epochs(raw, events=events, tmin=0., tmax=2.0,
                            baseline=None, preload=True, verbose=False)
        epochs.drop_bad(reject=dict(eeg=200e-6), verbose=False)
    psds, freqs = _epochs_psd_and_freqs(epochs, fmin=1., fmax=30.)
    psd_mean = psds.mean(axis=0)
    def integ(s, lo, hi):
        m = (freqs>=lo)&(freqs<hi)
        return float(np.trapz(s[m], freqs[m]))
    rows=[]
    for i,ch in enumerate(epochs.ch_names):
        s = psd_mean[i]
        total = integ(s, *TOTAL)
        band_abs = {b: integ(s, lo,hi) for b,(lo,hi) in BANDS.items()}
        rel = {b: (band_abs[b]/total if total>0 else np.nan) for b in BANDS}
        a1 = integ(s, *ALPHA1); a2 = integ(s, *ALPHA2)
        rows.append({"EDF_Channel": ch, "Cuban_Channel": normalize_label(ch),
                     **{f"Abs_{b}": band_abs[b] for b in BANDS},
                     **{f"Rel_{b}": rel[b] for b in BANDS},
                     "Abs_Alpha1": a1, "Abs_Alpha2": a2})
    df = pd.DataFrame(rows)
    df["Abs_Alpha"] = df["Abs_Alpha1"] + df["Abs_Alpha2"]
    return df, len(epochs)

# -------------------- Z-score helper --------------------
def add_zrel_zabs(df: pd.DataFrame, norms_age: pd.DataFrame) -> pd.DataFrame:
    def z_rel(p,m,s): return (np.log10(p)-m)/s if (p>0 and np.isfinite(p)) else np.nan
    def z_abs(v,m,s): return (np.log10(v)+12.0-m)/s if (v>0 and np.isfinite(v)) else np.nan
    out = df.copy()
    for b in BANDS:
        zr, za = [], []
        for _,r in out.iterrows():
            ch = r["Cuban_Channel"]
            nr = norms_age[(norms_age["Channel"]==ch)&(norms_age["Band"]==b)]
            if nr.empty:
                zr.append(np.nan); za.append(np.nan)
            else:
                nr = nr.iloc[0]
                zr.append(z_rel(r[f"Rel_{b}"], nr["RelMean"], nr["RelSD"]))
                za.append(z_abs(r[f"Abs_{b}"], nr["AbsMean_log10"], nr["AbsSD_log10"]))
        out[f"Zrel_{b}"]=zr; out[f"Zabs_{b}"]=za
    return out

# -------------------- Plotting --------------------
def _topomap_png(df: pd.DataFrame, key: str, title: str) -> bytes:
    xs, ys, zs = [], [], []
    for ch,(x,y) in POS.items():
        row = df[df["Cuban_Channel"]==ch]
        if row.empty: continue
        v = row.iloc[0].get(key, np.nan)
        if np.isfinite(v):
            xs.append(x); ys.append(y); zs.append(float(v))
    fig = plt.figure(figsize=(4.6,3.9))
    ax = fig.add_subplot(111, aspect="equal")
    if len(xs)==0:
        ax.text(0.5,0.5,"No data", ha="center", va="center")
    else:
        xs,ys,zs = np.array(xs), np.array(ys), np.array(zs)
        grid_x, grid_y = np.mgrid[-1:1:220j, -1:1:220j]
        pts = np.vstack([xs,ys]).T
        grid_z = griddata(pts, zs, (grid_x, grid_y), method="cubic")
        r = np.sqrt(grid_x**2 + grid_y**2); grid_z[r>1]=np.nan
        if key.startswith("Z"):
            levels = np.linspace(-3,3,13)
            cf = ax.contourf(grid_x, grid_y, grid_z, levels=levels, cmap="RdBu_r", vmin=-3, vmax=3)
        else:
            cf = ax.contourf(grid_x, grid_y, grid_z, levels=12)
        plt.colorbar(cf, ax=ax, shrink=0.75)
    ax.add_patch(Circle((0,0),1.0, fill=False, lw=2))
    ax.scatter([p[0] for p in POS.values()],[p[1] for p in POS.values()], c="k", s=12)
    ax.set_xticks([]); ax.set_yticks([]); ax.set_title(title)
    bio=io.BytesIO(); fig.savefig(bio, format="png", dpi=160, bbox_inches="tight"); plt.close(fig)
    return bio.getvalue()

def make_two_per_row_section(doc: Document, items: list):
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(0, len(items), 2):
        row_cells = table.add_row().cells
        for j in range(2):
            idx = i + j
            if idx >= len(items): break
            cap, png = items[idx]
            cell = row_cells[j]
            para = cell.paragraphs[0]; run = para.add_run()
            run.add_picture(io.BytesIO(png), width=Inches(2.8))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cap_p = cell.add_paragraph(cap); cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

# -------------------- Utilities --------------------
def get(df, ch, key):
    r = df[df["Cuban_Channel"]==ch]
    return float(r.iloc[0][key]) if not r.empty and key in r else np.nan

# -------------------- Diagnostics helpers --------------------
def compute_faa(dfz: pd.DataFrame):
    f3 = get(dfz,"F3","Rel_Alpha"); f4 = get(dfz,"F4","Rel_Alpha")
    if not (np.isfinite(f3) and np.isfinite(f4)) or (f3+f4)==0:
        return np.nan, "Insufficient data"
    faa_idx = (f3 - f4)/(f3 + f4) * 200
    if faa_idx < -20:
        tag = "Depression-leaning (FAA: left > right alpha)"
    elif faa_idx > 20:
        tag = "Right-dominant alpha (non-depressed pattern)"
    else:
        tag = "No strong FAA bias"
    return float(faa_idx), tag

def bipolar_theta_signature(dfz: pd.DataFrame):
    sites = {ch: get(dfz, ch, "Zrel_Theta") for ch in ["T8","P8","P4","O2"]}
    hits  = {ch: (np.isfinite(v) and v>=1.0) for ch,v in sites.items()}
    count = sum(1 for h in hits.values() if h)
    return sites, hits, (count>=2)

def region_alpha_reduced(dfz: pd.DataFrame, chans: list):
    vals = [get(dfz,ch,"Zrel_Alpha") for ch in chans]
    vals = [v for v in vals if np.isfinite(v)]
    if not vals: return False, np.nan, 0
    mean = float(np.nanmean(vals))
    n_low = sum(1 for v in vals if v <= -1.0)
    reduced = (mean <= -0.5) or (n_low >= 2)
    return reduced, mean, n_low

def bipolar_alpha_signature(dfz: pd.DataFrame):
    cen_red, cen_mean, cen_n = region_alpha_reduced(dfz, REGIONS["Central"])
    par_red, par_mean, par_n = region_alpha_reduced(dfz, REGIONS["Parietal"])
    occ_red, occ_mean, occ_n = region_alpha_reduced(dfz, REGIONS["Occipital"])
    alpha_ok = (cen_red and par_red) or (par_red and occ_red)
    details = {
        "Central":  {"reduced": cen_red, "mean": cen_mean, "sites_le_-1": cen_n},
        "Parietal": {"reduced": par_red, "mean": par_mean, "sites_le_-1": par_n},
        "Occipital":{"reduced": occ_red, "mean": occ_mean, "sites_le_-1": occ_n},
    }
    return details, alpha_ok

# -------------------- Main --------------------
def main():
    st.title("qEEG Report: Depression vs Bipolar (Streamlit Version)")
    st.markdown("Upload your EDF file and enter patient info. The Cuban norms CSV is embedded and loaded automatically.")

    edf_file = st.file_uploader("Select EDF file", type=["edf"])
    age = st.number_input("Patient Age", min_value=1, max_value=120, value=30, step=1)
    p2p = st.number_input("Artifact threshold (µV)", min_value=50.0, max_value=500.0, value=150.0, step=1.0)
    out_docx_name = st.text_input("Output Word report file name (optional)", value="qEEG_Report.docx")
    run_btn = st.button("Run Analysis")

    # Option 1: Read norms CSV from project directory
    try:
        norms_all = pd.read_csv("norms.csv")
    except Exception as e:
        st.error(f"Could not read norms.csv: {e}")
        return

    st.write("Cuban Norms CSV Preview:")
    st.dataframe(norms_all.head())

    if run_btn and edf_file and age:
        try:
            edf_path = f"temp_{edf_file.name}"
            with open(edf_path, "wb") as f: f.write(edf_file.getbuffer())

            raw0 = mne.io.read_raw_edf(edf_path, preload=False, verbose=False)
            unmatched = check_channel_mapping(raw0, norms_all)

            df_raw = compute_raw_powers(edf_path)
            df_clean, kept = compute_clean_powers(edf_path, p2p)
            norms_age = norms_all[norms_all["Age"]==age].copy()
            if norms_age.empty:
                st.error(f"No norms found for Age={age}.")
                return
            dfz = add_zrel_zabs(df_clean, norms_age)

            st.header("Summary")
            st.write(f"EDF: {edf_file.name}")
            st.write(f"Norms: norms.csv | Age: {age}")
            st.write(f"Artifact rejection: 2-s epochs, drop p2p > {p2p:.0f} µV. Kept epochs: {kept}")

            if unmatched:
                st.warning("Channels not found in Cuban norms (skipped in Z-maps):")
                for orig, mapped in unmatched:
                    st.write(f"EDF: {orig} -> Normalized: {mapped}")
            else:
                st.success("All EDF channels matched Cuban norms.")

            st.subheader("RAW Relative Maps")
            for b in ["Delta","Theta","Alpha","Beta"]:
                png = _topomap_png(df_raw, f"Rel_{b}", f"RAW Relative – {b}")
                st.image(png, caption=f"RAW Relative – {b}", use_column_width=True)

            st.subheader("Relative Z Maps (Cuban, CLEAN)")
            for b in ["Delta","Theta","Alpha","Beta"]:
                png = _topomap_png(dfz, f"Zrel_{b}", f"Z – {b} (Age {age})")
                st.image(png, caption=f"Z (Cuban) – {b}", use_column_width=True)

            st.header("Research-Based Diagnostics")

            faa_idx, faa_tag = compute_faa(dfz)
            faa_points = 2 if (np.isfinite(faa_idx) and faa_idx < -20) else 0
            st.write(f"FAA index (F3/F4): {faa_idx:.1f} -> {faa_tag} | Points={faa_points}")

            p3_alpha = get(dfz, "P3", "Rel_Alpha")
            p4_alpha = get(dfz, "P4", "Rel_Alpha")
            parietal_alpha_points = 0
            if np.isfinite(p3_alpha) and np.isfinite(p4_alpha):
                diff = p3_alpha - p4_alpha
                if diff > 0.1:
                    parietal_alpha_points = 1
                    asym_txt = "Left-dominant (Depression marker)"
                elif diff < -0.1:
                    asym_txt = "Right-dominant"
                else:
                    asym_txt = "Neutral"
            else:
                diff = np.nan
                asym_txt = "Insufficient data"
            st.write(f"Parietal Alpha Asymmetry (P3-P4): {diff:.3f} -> {asym_txt} | Points={parietal_alpha_points}")

            cz_pick = None
            for c in raw0.ch_names:
                if normalize_label(c) == "Cz":
                    cz_pick = c
                    break
            if cz_pick:
                psd_cz, freqs_cz = _raw_psd_and_freqs(raw0, cz_pick, 8., 12.)
                apf_hz = freqs_cz[np.argmax(psd_cz)] if len(freqs_cz) else np.nan
                apf_points = 3 if (np.isfinite(apf_hz) and apf_hz <= 9.5) else 0
                apf_txt = "Slowed (Depression marker)" if apf_points else "Normal"
                st.write(f"APF at Cz: {apf_hz:.2f} Hz -> {apf_txt} | Points={apf_points}")
            else:
                apf_hz, apf_points = np.nan, 0
                st.write("APF: Cz channel not found -> Skipped | Points=0")

            alpha_details, alpha_ok = bipolar_alpha_signature(dfz)
            alpha_reduction_points = 3 if alpha_ok else 0
            for r,info in alpha_details.items():
                mean_txt = f"{info['mean']:.2f}" if np.isfinite(info['mean']) else "NA"
                st.write(f"Alpha {r}: mean Z={mean_txt}, sites <= -1.0 = {info['sites_le_-1']} -> {'Reduced' if info['reduced'] else 'Not reduced'}")
            st.write(f"Alpha reduction rule: {'MET' if alpha_ok else 'NOT MET'} | Points={alpha_reduction_points}")

            sites, hits, theta_ok = bipolar_theta_signature(dfz)
            theta_points = 2 if theta_ok else 0
            theta_txt = ", ".join([f"{k}={sites[k]:.2f}" if np.isfinite(sites[k]) else f"{k}=NA" for k in ["T8","P8","P4","O2"]])
            st.write(f"Theta Z (T8,P8,P4,O2): {theta_txt} | >=2 sites >=+1.0 -> {'MET' if theta_ok else 'NOT MET'} | Points={theta_points}")

            p3_beta = get(dfz, "P3", "Zrel_Beta")
            p4_beta = get(dfz, "P4", "Zrel_Beta")
            parietal_beta_points = 1 if (np.isfinite(p3_beta) and np.isfinite(p4_beta) and (p3_beta >= 1.0 and p4_beta >= 1.0)) else 0
            beta_txt = "Bilateral elevation (Bipolar marker)" if parietal_beta_points else "Not elevated"
            st.write(f"Parietal Beta Bilateral (P3,P4): {p3_beta:.2f}, {p4_beta:.2f} -> {beta_txt} | Points={parietal_beta_points}")

            dep_score = faa_points + parietal_alpha_points + apf_points
            bp_score  = alpha_reduction_points + theta_points + parietal_beta_points
            total = dep_score + bp_score
            if total > 0:
                p_dep = dep_score / total
                p_bp  = bp_score  / total
            else:
                p_dep = p_bp = 0.5

            st.header("Probability Scores")
            st.write(f"P(Depression) = {p_dep*100:.1f}%")
            st.write(f"P(Bipolar)    = {p_bp*100:.1f}%")

            st.header("Scoring Breakdown")
            st.write(f"FAA: {faa_points}/2")
            st.write(f"Parietal Alpha Asymmetry: {parietal_alpha_points}/1")
            st.write(f"APF Slowing: {apf_points}/3")
            st.write(f"Alpha Reduction: {alpha_reduction_points}/3")
            st.write(f"Theta Right-Lower: {theta_points}/2")
            st.write(f"Parietal Beta Bilateral: {parietal_beta_points}/1")

            if dep_score > bp_score:
                final_call = "Depression-leaning"
            elif bp_score > dep_score:
                final_call = "Bipolar-leaning"
            else:
                final_call = "Ambiguous"
            st.header(f"Final: {final_call}")

            # Build Word report and allow download
            doc = Document()
            doc.add_heading("qEEG Report (Research-based Dx)", level=0)
            doc.add_paragraph(f"EDF: {os.path.basename(edf_path)}")
            doc.add_paragraph(f"Norms: norms.csv  |  Age: {age}")
            doc.add_paragraph(f"Artifact rejection: 2-s epochs, drop p2p > {p2p:.0f} µV. Kept epochs: {kept}")
            doc.add_paragraph("Band definitions: Delta 1–4, Theta 4–8, Alpha 8–12 (Alpha1 8–10, Alpha2 10–12), Beta 12–30 Hz.")

            if unmatched:
                doc.add_paragraph("Channels not found in Cuban norms (skipped in Z-maps):")
                for orig, mapped in unmatched:
                    doc.add_paragraph(f"  EDF: {orig}  ->  Normalized: {mapped}")
            else:
                doc.add_paragraph("All EDF channels matched Cuban norms.")

            doc.add_heading("RAW Relative Maps", level=1)
            raw_items = [(f"RAW Relative – {b}", _topomap_png(df_raw, f"Rel_{b}", f"RAW Relative – {b}"))
                         for b in ["Delta","Theta","Alpha","Beta"]]
            make_two_per_row_section(doc, raw_items)

            doc.add_heading("Relative Z Maps (Cuban, CLEAN)", level=1)
            z_items = [(f"Z (Cuban) – {b}", _topomap_png(dfz, f"Zrel_{b}", f"Z – {b} (Age {age})"))
                       for b in ["Delta","Theta","Alpha","Beta"]]
            make_two_per_row_section(doc, z_items)

            doc.add_heading("Research-Based Diagnostics", level=1)
            doc.add_paragraph(f"FAA index (F3/F4): {faa_idx:.1f}  ->  {faa_tag}  |  Points={faa_points}")
            doc.add_paragraph(f"Parietal Alpha Asymmetry (P3-P4): {diff:.3f} -> {asym_txt}  |  Points={parietal_alpha_points}")
            if cz_pick:
                doc.add_paragraph(f"APF at Cz: {apf_hz:.2f} Hz -> {apf_txt}  |  Points={apf_points}")
            else:
                doc.add_paragraph("APF: Cz channel not found -> Skipped  |  Points=0")
            for r,info in alpha_details.items():
                mean_txt = f"{info['mean']:.2f}" if np.isfinite(info['mean']) else "NA"
                doc.add_paragraph(f"Alpha {r}: mean Z={mean_txt}, sites <= -1.0 = {info['sites_le_-1']}  -> {'Reduced' if info['reduced'] else 'Not reduced'}")
            doc.add_paragraph(f"Alpha reduction rule: {'MET' if alpha_ok else 'NOT MET'}  |  Points={alpha_reduction_points}")
            doc.add_paragraph(f"Theta Z (T8,P8,P4,O2): {theta_txt}  |  >=2 sites >=+1.0 -> {'MET' if theta_ok else 'NOT MET'}  |  Points={theta_points}")
            doc.add_paragraph(f"Parietal Beta Bilateral (P3,P4): {p3_beta:.2f}, {p4_beta:.2f} -> {beta_txt}  |  Points={parietal_beta_points}")

            doc.add_heading("Probability Scores", level=2)
            doc.add_paragraph(f"P(Depression) = {p_dep*100:.1f}%")
            doc.add_paragraph(f"P(Bipolar)    = {p_bp*100:.1f}%")

            doc.add_heading("Scoring Breakdown", level=2)
            doc.add_paragraph(f"FAA: {faa_points}/2")
            doc.add_paragraph(f"Parietal Alpha Asymmetry: {parietal_alpha_points}/1")
            doc.add_paragraph(f"APF Slowing: {apf_points}/3")
            doc.add_paragraph(f"Alpha Reduction: {alpha_reduction_points}/3")
            doc.add_paragraph(f"Theta Right-Lower: {theta_points}/2")
            doc.add_paragraph(f"Parietal Beta Bilateral: {parietal_beta_points}/1")

            doc.add_heading(f"Final: {final_call}", level=2)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.download_button("Download Word Report", doc_io.getvalue(), file_name=out_docx_name or "qEEG_Report.docx")

        except Exception as e:
            st.error(str(e))

if __name__ == "__main__":
    main()
