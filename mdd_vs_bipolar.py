# mdd_bp_app.py
import streamlit as st

st.title("qEEG Apps1")

# نمایش متن گزارش فارسی در بالای اپلیکیشن
st.markdown("""
### پایگاه داده EEG هنجاری کوبا - توضیح بالینی

**مقدمه**  
پایگاه داده هنجاری EEG کوبا یکی از پرکاربردترین منابع بین المللی برای الکتروانسفالوگرافی کمی (qEEG) است. این پایگاه داده که از طریق مطالعات در مقیاس بزرگ در کوبا توسعه یافته و در چندین کشور تأیید شده است، مقادیر هنجاری طبقه بندی شده سنی قدرت طیفی EEG را در 10 تا 20 سایت الکترود استاندارد ارائه می دهد. این به عنوان یک نقطه مرجع طراحی شده است و به پزشکان و محققان اجازه می دهد تا فعالیت EEG یک بیمار را با هنجارهای معمول جمعیت مقایسه کنند.

**نشانگرهای EEG و معنای بالینی آنها**  
**فرکانس پیک آلفا (APF)**  
APF فرکانس غالب در باند آلفا (8-12 هرتز) است. از نظر هنجاری، APF در دوران کودکی افزایش می یابد و در بزرگسالی تثبیت می شود. کاهش سرعت APF (≤9.5 هرتز در بزرگسالان) با اختلال افسردگی اساسی و زوال شناختی مرتبط است.

**عدم تقارن آلفا پیشانی (FAA)**  
FAA به عنوان تفاوت نسبی آلفای بین سایت های پیشانی چپ (F3) و راست (F4) محاسبه می شود. FAA منفی (آلفای چپ بزرگتر) با افسردگی همراه است، در حالی که آلفای غالب راست نشان دهنده انعطاف پذیری در برابر حالت های افسردگی است.

**عدم تقارن آلفای جداری**  
تفاوت بین قدرت آلفای جداری چپ (P3) و راست (P4) نشانگر دیگری برای افسردگی فراهم می کند. آلفای غالب چپ در بیماران افسردگی مشاهده شده است.

**تتا اضافی (زنجیره راست پایین)**  
افزایش توان تتا (Z ≥ +1.0) در حداقل دو الکترود زنجیره پایین راست (T8، P8، P4، O2) نشانگر اختلال دوقطبی در نظر گرفته می شود. این نشان دهنده فعالیت موج آهسته غیرمعمول در مناطق گیجگاهی-جداری است.

**کاهش آلفا (مرکزی-جداری-اکسیپیتال)**  
کاهش قدرت آلفا در نواحی مرکزی و جداری یا نواحی جداری و پس سری یکی دیگر از شاخص های اختلال دوقطبی است. این الگو نشان می دهد کمبود تولید ریتم آلفا خلفی.

**ارتفاع دو طرفه بتا جداری**  
ارتفاع متقارن دو طرفه بتا در P3 و P4 (Z ≥+1.0) با اختلال دوقطبی همراه است. افزایش فعالیت بتا ممکن است منعکس کننده تحریک پذیری قشر مغز باشد.

**سیستم امتیازدهی وزنی**  
برای طبقه بندی مبتنی بر تحقیق بین اختلال افسردگی اساسی (MDD) و اختلال دوقطبی (BD)، نشانگرها بر اساس ارزش تشخیصی آنها نقاط وزنی اختصاص داده می شوند:  
• FAA (عدم تقارن آلفای پیشانی): تا 2 نقطه (نشانگر افسردگی)  
• عدم تقارن آلفای جداری: تا 1 نقطه (نشانگر افسردگی)  
• کاهش APF در Cz: تا 3 نقطه (نشانگر افسردگی)  
• کاهش آلفا (مرکزی/جداری / پس سری): تا 3 امتیاز (نشانگر دو قطبی)  
• بیش از حد تتا در زنجیره پایین سمت راست: تا 2 امتیاز (نشانگر دوقطبی)  
• ارتفاع دو طرفه بتا جداری: تا 1 نقطه (نشانگر دوقطبی)

نمره افسردگی از مجموع FAA ، عدم تقارن آلفای جداری و کاهش سرعت APF (حداکثر = 6) محاسبه می شود. نمره دوقطبی از کاهش آلفا، بیش از حد تتا و ارتفاع دو طرفه بتا جداری (حداکثر = 6) محاسبه می شود. احتمالات نهایی با نرمال کردن این امتیازات به درصد به دست می آیند و P(Depression) و P(Bi-قطبی) به دست می آیند.

**نتیجه**  
بانک اطلاعاتی هنجاری کوبا پزشکان را قادر می سازد تا یافته های qEEG را در چارچوب استانداردهای مبتنی بر جمعیت تفسیر کنند. با استفاده از نمرات Z و سیستم های امتیازدهی ساختاریافته، نشانگرهای عینی برای تمایز بین افسردگی و اختلال دوقطبی و همچنین ارزیابی گسترده تر اعصاب و روان فراهم می کند.
""")

# ... rest of your app code ...


import os
import io
import time
import tempfile
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document

import mdd_vs_bipolar as mddbp  # our analysis helpers (updated file)

# -----------------------
# Config
# -----------------------
st.set_page_config(page_title="qEEG: MDD vs Bipolar", layout="wide")
NORM_PATH = "cuban_REC_RD_norms_EC_5_87.csv"   # keep this CSV in the same folder

BANDS = ["Delta", "Theta", "Alpha", "Beta"]

# -----------------------
# Helpers
# -----------------------
@st.cache_data(show_spinner=False)
def load_norms_df(path: str) -> pd.DataFrame:
    df = pd.read_csv(path)
    # expected columns: Age, Channel, Band, AbsMean_log10, AbsSD_log10, RelMean, RelSD
    # clean up any stray spaces/case
    df.columns = [c.strip() for c in df.columns]
    df["Band"] = df["Band"].astype(str)
    return df

def get_norms_for_age(norms_all: pd.DataFrame, target_age: int) -> pd.DataFrame:
    # pick the closest available age in the norms
    ages = norms_all["Age"].astype(int)
    closest_age = int(ages.iloc[(ages - target_age).abs().argmin()])
    return norms_all[norms_all["Age"].astype(int) == closest_age].copy(), closest_age

@st.cache_data(show_spinner=True)
def save_uploaded_file_to_temp(uploaded_file) -> str:
    # Streamlit gives an UploadedFile; we persist it to a temp EDF path
    suffix = os.path.splitext(uploaded_file.name)[-1] or ".edf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        return tmp.name

# convenience to show two columns of images
def show_maps_grid(items, cols=2):
    if not items:
        st.info("No maps to display.")
        return
    rows = (len(items) + cols - 1) // cols
    idx = 0
    for _ in range(rows):
        cols_list = st.columns(cols, gap="large")
        for c in range(cols):
            if idx >= len(items):
                break
            title, img_bytes = items[idx]
            with cols_list[c]:
                st.image(img_bytes, caption=title, use_container_width=True)
            idx += 1

# -----------------------
# UI
# -----------------------
st.title("qEEG – MDD vs Bipolar (Demo)")
st.caption("Upload an EDF, enter age, and generate raw/Z maps, markers, and a Word report.")

left, right = st.columns([1, 1])

with left:
    edf_up = st.file_uploader("Upload EEG (EDF)", type=["edf"])
    age_in = st.number_input("Patient age", min_value=5, max_value=87, value=25, step=1)
    artifact_thresh = st.number_input("Artifact peak-to-peak threshold (µV)", min_value=50, max_value=1000, value=150, step=10)
    use_clean = st.checkbox("Use simple artifact rejection (2s epochs)", value=True)
    run_btn = st.button("Run analysis", type="primary")

with right:
    st.markdown("**Norms file**")
    st.write(f"Using: `{NORM_PATH}` (place it in the app folder)")
    st.divider()

# Load norms (once)
try:
    norms_all = load_norms_df(NORM_PATH)
    norms_ok = True
except Exception as e:
    norms_ok = False
    st.error(f"Could not load norms CSV `{NORM_PATH}`: {e}")

# -----------------------
# Analysis
# -----------------------
if run_btn:
    if not edf_up:
        st.warning("Please upload an EDF file first.")
        st.stop()
    if not norms_ok:
        st.stop()

    # Save EDF to a temp path
    with st.status("Saving uploaded EDF...", expanded=False):
        edf_path = save_uploaded_file_to_temp(edf_up)

    # Pick norms for nearest age
    norms_age, used_age = get_norms_for_age(norms_all, int(age_in))

    # Compute powers
    with st.status("Computing band powers...", expanded=True) as s:
        t0 = time.time()
        if use_clean:
            df_clean, kept = mddbp.compute_clean_powers(edf_path, p2p_thresh=int(artifact_thresh))
            df_raw = df_clean.copy()
            s.update(label=f"Computed clean powers from {kept} accepted epochs.")
        else:
            df_raw = mddbp.compute_raw_powers(edf_path)
            s.update(label=f"Computed raw powers.")
        s.update(state="complete", expanded=False)

    st.success(f"EEG channels found: {len(df_raw.index)} | Norms age used: {used_age}y | Elapsed {time.time()-t0:.1f}s")

    # Z-scores vs Cuban norms
    with st.status("Computing Z-scores (Cuban norms)...", expanded=False):
        dfz = mddbp.add_zrel_zabs(df_raw, norms_age)

    # FAA & Markers & Score
    faa_val, faa_tag = mddbp.compute_faa(dfz)
    alpha_details, alpha_ok = mddbp.bipolar_alpha_signature(dfz)
    theta_sites, theta_hits, theta_ok = mddbp.bipolar_theta_signature(dfz)
    probs = mddbp.compute_diagnostic_score(dfz)

    # -----------------------
    # Visuals (topomaps)
    # -----------------------
    st.subheader("Topomaps")

    raw_items = [(f"RAW – {b}", mddbp._topomap_png(df_raw, f"Rel_{b}", f"RAW – {b}")) for b in BANDS]
    z_items = [(f"Z – {b}", mddbp._topomap_png(dfz, f"Zrel_{b}", f"Z – {b}")) for b in BANDS]

    st.markdown("**Raw Relative Power**")
    show_maps_grid(raw_items)

    st.markdown("**Z-scores vs Cuban norms**")
    show_maps_grid(z_items)

    # -----------------------
    # Results summary
    # -----------------------
    st.subheader("Markers & Summary")

    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("FAA", f"{faa_val:.1f}" if np.isfinite(faa_val) else "N/A", help=f"Frontal Alpha Asymmetry – {faa_tag}")
        st.caption(f"Interpretation: **{faa_tag}**")

    with c2:
        st.write("**Bipolar Alpha Reduction**")
        st.write(f"- Central reduced:  {alpha_details['Central']['reduced']} (mean {alpha_details['Central']['mean']:.2f})")
        st.write(f"- Parietal reduced: {alpha_details['Parietal']['reduced']} (mean {alpha_details['Parietal']['mean']:.2f})")
        st.write(f"- Occipital reduced:{alpha_details['Occipital']['reduced']} (mean {alpha_details['Occipital']['mean']:.2f})")
        st.success("Alpha reduction signature: **Present**" if alpha_ok else "Alpha reduction signature: **Absent**")

    with c3:
        st.write("**Right-lower Theta Chain (T8,P8,P4,O2)**")
        st.write({k: f"{v:.2f}" for k, v in theta_sites.items()})
        st.success("Theta chain: **Present**" if theta_ok else "Theta chain: **Absent**")

    st.divider()
    st.subheader("Depression vs Bipolar – Probabilities")
    st.write({k: f"{v*100:.1f}%" for k, v in probs.items()})
    final_call = "Ambiguous"
    if probs["Bipolar"] >= 0.6:
        final_call = "Bipolar-leaning"
    elif probs["Depression"] >= 0.6:
        final_call = "Depression-leaning"
    st.info(f"**Final call:** {final_call}")

    # -----------------------
    # Build Word report
    # -----------------------
    st.subheader("Export Word Report")
    doc = Document()
    doc.add_heading("qEEG – MDD vs Bipolar Report", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Age: {age_in}  |  Norms age used: {used_age}\n")
    p.add_run(f"Artifact threshold: {artifact_thresh} µV  |  Cleaned: {use_clean}\n")

    doc.add_heading("Raw Relative Power", level=2)
    mddbp.make_two_per_row_section(doc, raw_items)

    doc.add_heading("Z-score Maps (Cuban norms)", level=2)
    mddbp.make_two_per_row_section(doc, z_items)

    doc.add_heading("Markers & Scores", level=2)
    t = doc.add_table(rows=0, cols=2)
    r = t.add_row().cells
    r[0].text = "FAA"
    r[1].text = f"{faa_val:.1f} ({faa_tag})" if np.isfinite(faa_val) else "N/A"

    r = t.add_row().cells
    r[0].text = "Alpha reduction signature"
    r[1].text = "Present" if alpha_ok else "Absent"

    r = t.add_row().cells
    r[0].text = "Theta chain"
    r[1].text = "Present" if theta_ok else "Absent"

    r = t.add_row().cells
    r[0].text = "Probabilities"
    r[1].text = f"Depression: {probs['Depression']*100:.1f}%   |   Bipolar: {probs['Bipolar']*100:.1f}%"

    doc.add_paragraph()
    doc.add_paragraph(f"Final call: {final_call}")

    # Save to bytes and present download button
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "Download Word Report",
        data=buf.getvalue(),
        file_name="qEEG_MDDvsBipolar_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.success("Done ✅  (If something looks empty, double-check: EDF had data, and norms CSV is in the folder.)")
